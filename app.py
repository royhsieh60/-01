import os
import sys
import uuid
import io
import time
import json
import re
import random
import urllib.parse
import threading
import concurrent.futures
import base64
import traceback
import requests
from PIL import Image
from pypdf import PdfReader
from fastapi import FastAPI, Request, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse

# 👑 引入 LINE Messaging API
from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError
from linebot.models import (
    MessageEvent, TextMessage, ImageMessage, FileMessage, 
    TextSendMessage, ImageSendMessage, QuickReply, QuickReplyButton, 
    MessageAction, TemplateSendMessage, ButtonsTemplate, CarouselTemplate, CarouselColumn
)

# AI 官方套件
from huggingface_hub import InferenceClient
from google import genai
from docx import Document
from docx.shared import Inches
import pymongo
import ssl
import certifi

# ==========================================
# 👑 基礎環境設定
# ==========================================
os.environ["SSL_CERT_FILE"] = certifi.where()
os.environ["REQUESTS_CA_BUNDLE"] = certifi.where()
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError: pass
else: ssl._create_default_https_context = _create_unverified_https_context

os.makedirs("static", exist_ok=True)
app = FastAPI()

@app.get("/static/{filename}")
async def serve_file(filename: str):
    file_path = os.path.join("static", filename)
    if os.path.exists(file_path): return FileResponse(file_path)
    raise HTTPException(status_code=404, detail="File not found")

LINE_CHANNEL_ACCESS_TOKEN = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN", "").strip()
LINE_CHANNEL_SECRET = os.environ.get("LINE_CHANNEL_SECRET", "").strip()
SPACE_URL = os.environ.get("SPACE_URL", "").strip()
MONGO_URI = os.environ.get("MONGO_URI", "").strip()

line_bot_api = LineBotApi(LINE_CHANNEL_ACCESS_TOKEN)
handler = WebhookHandler(LINE_CHANNEL_SECRET)

try:
    mongo_client = pymongo.MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000, tlsAllowInvalidCertificates=True)
    db = mongo_client["LineBotDB"]
    chat_collection = db["chat_history"]
    prefs_collection = db["user_prefs"]
except Exception as e: print("資料庫連線失敗: " + str(e))

user_last_prompt = {}

def save_error_log(error_text):
    doc_name = "ErrorLog_" + uuid.uuid4().hex[:8] + ".txt"
    doc_path = os.path.join("static", doc_name)
    with open(doc_path, "w", encoding="utf-8") as f:
        f.write("========== 原始錯誤日誌 ==========\n")
        f.write(time.strftime("%Y-%m-%d %H:%M:%S") + "\n\n")
        f.write(str(error_text))
    return SPACE_URL + "/static/" + doc_name + "?openExternalBrowser=1"

def show_loading_animation(user_id):
    try: requests.post("https://api.line.me/v2/bot/chat/loading/start", headers={"Content-Type": "application/json", "Authorization": "Bearer " + LINE_CHANNEL_ACCESS_TOKEN}, json={"chatId": user_id, "loadingSeconds": 20})
    except: pass

def clean_line_text(text: str) -> str:
    if not text: return ""
    text = re.sub(r'#+\s+', '', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    return text.strip()

def clean_math_fallback(text: str) -> str:
    if not text: return ""
    def parse_matrix(match):
        content = match.group(1)
        rows = content.split(r'\\')
        out = []
        for r in rows:
            cols = [c.strip() for c in r.split('&')]
            out.append("[ " + " | ".join(cols) + " ]")
        return "\n" + "\n".join(out) + "\n"
    text = re.sub(r'\\begin\{[bpBvV]?matrix\}(.*?)\\end\{[bpBvV]?matrix\}', parse_matrix, text, flags=re.DOTALL)
    while True:
        new_text = re.sub(r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'(\1)/(\2)', text)
        if new_text == text: break
        text = new_text
    replacements = {r'\times': '×', r'\cdot': '·', r'\div': '÷', r'\sqrt': '√', r'\pi': 'π', r'\theta': 'θ', r'\pm': '±', r'\infty': '∞', r'\approx': '≈', r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\{': '{', r'\}': '}', r'\$': '', r'&': ' ', r'\boldsymbol': '', r'\mathbf': '', r'\left': '', r'\right': '', r'\text': '', r'\quad': ' ', r'\[': '', r'\]': '', r'\(': '', r'\)': '', r'\\': '\n'}
    for k, v in replacements.items(): text = text.replace(k, v)
    return re.sub(r'\\[a-zA-Z]+', '', text).strip()

def process_inline_math(text: str) -> str:
    if not text: return ""
    text = re.sub(r'\\\((.*?)\\\)', lambda m: clean_math_fallback(m.group(1)), text, flags=re.DOTALL)
    return re.sub(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', lambda m: clean_math_fallback(m.group(1)), text, flags=re.DOTALL)

def clean_history_text(text: str) -> str:
    if not text: return ""
    prefixes = [r"🌟 【雙開整合解答】\n\n", r"🟢 【(.*?) 獨立解答】(.*?)\n\n", r"🔵 【(.*?) 獨立解答】(.*?)\n\n", r"🔄 \[(.*?)\]\n", r"⚠️ \[(.*?)\]\n\n", r"❌ \[(.*?)\]\n\n"]
    for p in prefixes: text = re.sub(p, "", text, flags=re.DOTALL)
    return text.strip()

# ==========================================
# 👑 官方 API 通道：Gemini (絕對不會被 IP 封鎖)
# ==========================================
raw_gemini = os.environ.get("GEMINI_API_KEY", "").replace("\n", ",").replace("\r", ",")
GEMINI_KEYS = [k.strip().strip('"').strip("'") for k in raw_gemini.split(",") if k.strip()]
if not GEMINI_KEYS: GEMINI_KEYS = ["dummy_key"]

current_gemini_key_index = 0
gemini_lock = threading.Lock()

def get_gemini_client():
    with gemini_lock: return genai.Client(api_key=GEMINI_KEYS[current_gemini_key_index])

def get_next_gemini_client():
    global current_gemini_key_index
    with gemini_lock:
        current_gemini_key_index = (current_gemini_key_index + 1) % len(GEMINI_KEYS)
        return genai.Client(api_key=GEMINI_KEYS[current_gemini_key_index])

def run_gemini(contents, model_name='gemini-1.5-flash', max_retries=len(GEMINI_KEYS)*2):
    client = get_gemini_client()
    for attempt in range(max_retries):
        try:
            res = client.models.generate_content(model=model_name, contents=contents)
            if not res or not res.text: raise ValueError(model_name + " 回傳空值")
            return res.text
        except Exception as e:
            error_str = str(e).lower()
            if "404" in error_str: raise ValueError("模型不存在 (404)")
            if "429" in error_str or "quota" in error_str or "api_key_invalid" in error_str:
                if len(GEMINI_KEYS) > 1:
                    client = get_next_gemini_client()
                    continue
                else: raise ValueError("Gemini 額度耗盡或無效: " + str(e)[:30])
            if attempt < max_retries - 1:
                time.sleep(1)
                continue
            raise e

# ==========================================
# 👑 官方 API 通道：Hugging Face (精準處理 DNS 與 402)
# ==========================================
raw_hf = os.environ.get("HF_TOKEN", "").replace("\n", ",").replace("\r", ",")
HF_KEYS = [k.strip().strip('"').strip("'") for k in raw_hf.split(",") if k.strip()]
if not HF_KEYS: HF_KEYS = ["dummy_key"]

current_hf_key_index = 0
hf_lock = threading.Lock()

def get_current_hf_key():
    with hf_lock: return HF_KEYS[current_hf_key_index]

def get_next_hf_key():
    global current_hf_key_index
    with hf_lock:
        current_hf_key_index = (current_hf_key_index + 1) % len(HF_KEYS)
        return HF_KEYS[current_hf_key_index]

def run_hf_model(model_id, messages):
    last_err = ""
    url = "https://api-inference.huggingface.co/models/" + model_id + "/v1/chat/completions"
    
    for attempt in range(len(HF_KEYS) * 2): 
        current_key = get_current_hf_key()
        if current_key == "dummy_key": raise ValueError("HF_TOKEN 未設定")
            
        headers = {"Authorization": "Bearer " + current_key, "Content-Type": "application/json"}
        payload = {"model": model_id, "messages": messages, "max_tokens": 2048, "temperature": 0.7}
        
        try:
            res = requests.post(url, headers=headers, json=payload, timeout=20)
            if res.status_code == 200:
                data = res.json()
                if "choices" in data and len(data["choices"]) > 0: return data["choices"][0]["message"]["content"]
                raise ValueError("HF 回傳異常")
            elif res.status_code == 402:
                raise ValueError("【402 官方付費牆】該模型需付費存取。")
            elif res.status_code in [401, 403, 429]:
                get_next_hf_key()
                last_err = "金鑰受限/額度耗盡 (" + str(res.status_code) + ")"
            elif res.status_code == 503:
                time.sleep(2)
                continue
            else:
                get_next_hf_key()
                last_err = "HTTP " + str(res.status_code)
        except requests.exceptions.ConnectionError:
            raise ValueError("【DNS 解析崩潰】HF 雲端主機網路中斷 (NameResolutionError)")
        except Exception as e:
            if "付費牆" in str(e) or "DNS" in str(e): raise e
            get_next_hf_key()
            last_err = str(e)[:30]
            
    raise ValueError("HF 模型失敗: " + last_err)

# ==========================================
# 👑 引擎調度 (安全降級網)
# ==========================================
TEXT_ENGINES = {
    "gemini-2.5-flash": "Gemini 2.5 Flash",
    "gemini-1.5-flash": "Gemini 1.5 Flash",
    "gemini-1.5-pro": "Gemini 1.5 Pro",
    "Qwen/Qwen2.5-72B-Instruct": "Qwen 72B",
    "Qwen/Qwen2.5-7B-Instruct": "Qwen 7B",
    "meta-llama/Llama-3.2-3B-Instruct": "Llama 3.2-3B",
    "merge": "合併雙開"
}

def smart_engine_runner(target_engine, hf_messages, gemini_contents):
    if target_engine.startswith("Qwen") or target_engine.startswith("meta-llama"):
        return run_hf_model(target_engine, hf_messages)
    else:
        return run_gemini(gemini_contents, model_name=target_engine if target_engine.startswith("gemini") else "gemini-1.5-flash")

# ==========================================
# 👑 生圖引擎 (完全依賴 HF 官方 API，防止被 Pollinations 封鎖 IP)
# ==========================================
ENGINE_MODELS = {
    "auto": "black-forest-labs/FLUX.1-schnell", "sdxl": "stabilityai/stable-diffusion-xl-base-1.0",
    "flux": "black-forest-labs/FLUX.1-schnell", "anime": "cagliostrolab/animagine-xl-3.1"
}
ENGINE_LABELS = {"auto": "智能", "sdxl": "SDXL(官方)", "flux": "Flux", "anime": "動漫風"}

def generate_image_and_save(engine, prompt, seed):
    model_id = ENGINE_MODELS.get(engine, ENGINE_MODELS["auto"])
    img, last_err = None, ""
    
    url = "https://api-inference.huggingface.co/models/" + model_id
    for attempt in range(len(HF_KEYS)):
        current_key = get_current_hf_key()
        if current_key == "dummy_key":
            get_next_hf_key()
            continue
        try:
            res = requests.post(url, headers={"Authorization": "Bearer " + current_key}, json={"inputs": prompt}, timeout=20)
            if res.status_code == 200:
                img = Image.open(io.BytesIO(res.content))
                break
            elif res.status_code == 402:
                raise ValueError("生圖模型已被官方加入付費牆 (402)。請改用本機部署。")
            else:
                last_err += "HF(" + str(res.status_code) + ") "
                get_next_hf_key()
        except requests.exceptions.ConnectionError:
            raise ValueError("HF 雲端主機 DNS 解析崩潰，生圖連線失敗。")
        except Exception as e: 
            if "付費牆" in str(e) or "DNS" in str(e): raise e
            last_err += "ERR(" + str(e)[:15] + ") "
            get_next_hf_key()

    if img is None: raise ValueError("生圖失敗: " + last_err)
    if img.mode != 'RGB': img = img.convert('RGB')
    name = uuid.uuid4().hex + ".jpg"
    path = os.path.join("static", name)
    img.save(path, format="JPEG", quality=90)
    return SPACE_URL + "/static/" + name

# ==========================================
# 👑 選單架構
# ==========================================
def get_quick_replies(user_id):
    try: pref = prefs_collection.find_one({"user_id": user_id}) or {}
    except: pref = {}
    bypass_mode = pref.get("bypass_mode", False)
    buttons = [
        QuickReplyButton(action=MessageAction(label="⚙️ 大腦設定", text="⚙️ 大腦設定")),
        QuickReplyButton(action=MessageAction(label="🎨 生圖設定", text="🎨 生圖設定")),
        QuickReplyButton(action=MessageAction(label="⚡直通生圖: 🟢開" if bypass_mode else "⚡直通生圖: 🔴關", text="/toggle_bypass")),
        QuickReplyButton(action=MessageAction(label="🩺 系統檢查", text="系統檢查")),
        QuickReplyButton(action=MessageAction(label="🧹 清除記憶", text="/clear"))
    ]
    return QuickReply(items=buttons)

def send_engine_family_menu(user_id, reply_token):
    col1 = CarouselColumn(title="👑 官方與免授權系列", text="選擇引擎家族以展開版本參數", actions=[
        MessageAction(label="Gemini 系列 (Google)", text="/engine_family gemini"),
        MessageAction(label="合併雙開 (最強)", text="/engine merge"),
        MessageAction(label="取消", text="OK")
    ])
    col2 = CarouselColumn(title="🤖 開源巨獸系列", text="選擇引擎家族以展開版本參數", actions=[
        MessageAction(label="Qwen 系列 (阿里)", text="/engine_family qwen"),
        MessageAction(label="Llama 系列 (Meta)", text="/engine_family llama"),
        MessageAction(label="返回", text="OK")
    ])
    send_line_messages(user_id, reply_token, [TemplateSendMessage(alt_text="大腦家族選單", template=CarouselTemplate(columns=[col1, col2]))])

def send_specific_engine_menu(user_id, reply_token, family):
    columns = []
    try: current = (prefs_collection.find_one({"user_id": user_id}) or {}).get("text_engine", "")
    except: current = ""
    def mark(eid): return "✅ " if current == eid else ""

    if family == "gemini":
        columns.append(CarouselColumn(title="Gemini 官方", text="穩定高速", actions=[
            MessageAction(label=mark("gemini-2.5-flash") + "2.5 Flash", text="/engine gemini-2.5-flash"),
            MessageAction(label=mark("gemini-1.5-flash") + "1.5 Flash (最穩)", text="/engine gemini-1.5-flash"),
            MessageAction(label=mark("gemini-1.5-pro") + "1.5 Pro (深度分析)", text="/engine gemini-1.5-pro")
        ]))
    elif family == "qwen":
        columns.append(CarouselColumn(title="Qwen 大參數", text="最強邏輯", actions=[
            MessageAction(label=mark("Qwen/Qwen2.5-72B-Instruct") + "72B (PRO)", text="/engine Qwen/Qwen2.5-72B-Instruct"),
            MessageAction(label=mark("Qwen/Qwen2.5-32B-Instruct") + "32B (PRO)", text="/engine Qwen/Qwen2.5-32B-Instruct"),
            MessageAction(label=mark("Qwen/Qwen2.5-7B-Instruct") + "7B (免費)", text="/engine Qwen/Qwen2.5-7B-Instruct")
        ]))
    elif family == "llama":
        columns.append(CarouselColumn(title="Llama 3", text="Meta 官方模型", actions=[
            MessageAction(label=mark("meta-llama/Llama-3.2-3B-Instruct") + "3.2 3B", text="/engine meta-llama/Llama-3.2-3B-Instruct"),
            MessageAction(label="回主分類", text="⚙️ 大腦設定"),
            MessageAction(label="取消", text="OK")
        ]))
    if columns: send_line_messages(user_id, reply_token, [TemplateSendMessage(alt_text="版本選單", template=CarouselTemplate(columns=columns))])

def send_image_engine_menu(user_id, reply_token):
    try: current = (prefs_collection.find_one({"user_id": user_id}) or {}).get("image_engine", "auto")
    except: current = "auto"
    col1 = CarouselColumn(title="🎨 生圖模型 (僅官方)", text="Hugging Face 原生通道", actions=[
        MessageAction(label=("✅ " if current=="auto" else "") + "FLUX.1-schnell", text="/image_engine auto"),
        MessageAction(label=("✅ " if current=="sdxl" else "") + "SDXL-1.0", text="/image_engine sdxl"),
        MessageAction(label=("✅ " if current=="anime" else "") + "Animagine-XL", text="/image_engine anime")
    ])
    send_line_messages(user_id, reply_token, [TemplateSendMessage(alt_text="生圖選單", template=CarouselTemplate(columns=[col1]))])

def get_batch_template():
    return TemplateSendMessage(alt_text="批量生成選單", template=ButtonsTemplate(text="✨ 圖像已生成！需要批量產出嗎？", actions=[MessageAction(label="🔄 再來 1 張", text="/batch 1"), MessageAction(label="🚀 批量 3 張", text="/batch 3"), MessageAction(label="🔥 批量 10 張", text="/batch 10")]))

def send_line_messages(user_id, reply_token, messages):
    if not messages: return
    messages[-1].quick_reply = get_quick_replies(user_id)
    for i in range(0, len(messages), 5):
        batch = messages[i:i+5]
        if i == 0 and reply_token:
            try: line_bot_api.reply_message(reply_token, batch)
            except:
                try: line_bot_api.push_message(user_id, batch)
                except Exception as push_e: print("[Push 失敗]: " + str(push_e))
        else:
            try: line_bot_api.push_message(user_id, batch)
            except Exception as push_e: print("[Push 失敗]: " + str(push_e))

# ==========================================
# 👑 系統排查報告
# ==========================================
def run_system_diagnostic_to_file():
    report = "🩺 【LineBot 系統排查報告】\n" + "="*50 + "\n\n[Hugging Face 引擎測試]\n"
    for i, k in enumerate(HF_KEYS):
        if not k or k == "dummy_key":
            report += "  ❌ Key [" + str(i) + "]: 未設定\n"
            continue
        try:
            res = requests.get("https://huggingface.co/api/whoami-v2", headers={"Authorization": "Bearer " + k}, timeout=5)
            if res.status_code == 200: report += "  ✅ Key [" + str(i) + "]: 授權成功\n"
            else: report += "  ⚠️ Key [" + str(i) + "]: 狀態 (" + str(res.status_code) + ")\n"
        except Exception as e: report += "  ❌ Key [" + str(i) + "]: 連線失敗 (" + str(e)[:30] + ")\n"
            
    report += "\n[Google Gemini 引擎測試]\n"
    for i, k in enumerate(GEMINI_KEYS):
        if not k or k == "dummy_key":
            report += "  ❌ Key [" + str(i) + "]: 未設定\n"
            continue
        try:
            c = genai.Client(api_key=k)
            c.models.generate_content(model='gemini-1.5-flash', contents="hi")
            report += "  ✅ Key [" + str(i) + "]: 生成成功\n"
        except Exception as e:
            report += "  ❌ Key [" + str(i) + "]: 測試失敗 (" + str(e)[:30] + ")\n"
        
    doc_name = "Health_Report_" + uuid.uuid4().hex[:6] + ".txt"
    path = os.path.join("static", doc_name)
    with open(path, "w", encoding="utf-8") as f: f.write(report)
    return SPACE_URL + "/static/" + doc_name + "?openExternalBrowser=1"

# ==========================================
# 👑 終極 DOCX 解析器
# ==========================================
def parse_and_build_content(text, user_id):
    if not text: text = ""
    text = process_inline_math(text)
    line_messages = []
    doc = Document()
    doc.add_heading('AI 深度解析報告', 0)

    text = re.sub(r'\\\[(.*?)\\\]', r'__MATH__\1__MATH_END__', text, flags=re.DOTALL)
    text = re.sub(r'\$\$(.*?)\$\$', r'__MATH__\1__MATH_END__', text, flags=re.DOTALL)
    parts = re.compile(r'(__CHART__.*?__CHART_END__|__GRAPH__.*?__GRAPH_END__|__IMAGE__.*?__IMAGE_END__|__MATH__.*?__MATH_END__)', re.DOTALL).split(text)

    current_text_block = ""
    for part in parts:
        if not part.strip(): continue
        img_url, is_math, raw_content = None, False, ""
        
        if part.startswith('__CHART__'):
            raw_content = part.replace('__CHART__', '').replace('__CHART_END__', '').strip()
            img_url = "https://quickchart.io/chart?c=" + urllib.parse.quote(raw_content)
        elif part.startswith('__GRAPH__'):
            raw_content = part.replace('__GRAPH__', '').replace('__GRAPH_END__', '').strip()
            img_url = "https://quickchart.io/graphviz?graph=" + urllib.parse.quote(raw_content)
        elif part.startswith('__MATH__'):
            raw_content = part.replace('__MATH__', '').replace('__MATH_END__', '').strip()
            img_url = "https://latex.codecogs.com/png.image?\\dpi{300}\\bg_white\\space " + urllib.parse.quote(raw_content)
            is_math = True
        elif part.startswith('__IMAGE__'):
            # 捨棄不穩定的外部影像搜尋，直接以文字顯示
            raw_content = part.replace('__IMAGE__', '').replace('__IMAGE_END__', '').strip()
            img_url = None
        else:
            current_text_block += part + "\n"
            continue
        
        if current_text_block.strip():
            cleaned = clean_line_text(current_text_block)
            if cleaned: doc.add_paragraph(cleaned)
            current_text_block = ""
            
        if img_url:
            img_success = False
            try:
                res = requests.get(img_url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
                if res.status_code == 200 and 'image' in res.headers.get('Content-Type', '').lower():
                    temp_img_path = "static/temp_" + uuid.uuid4().hex + ".jpg"
                    with open(temp_img_path, 'wb') as f: f.write(res.content)
                    try:
                        target_width = None
                        with Image.open(temp_img_path) as tmp_img:
                            if tmp_img.mode in ('RGBA', 'LA') or (tmp_img.mode == 'P' and 'transparency' in tmp_img.info):
                                alpha = tmp_img.convert('RGBA').split()[-1]
                                bg = Image.new("RGB", tmp_img.size, (255, 255, 255))
                                bg.paste(tmp_img, mask=alpha)
                                bg.save(temp_img_path, 'JPEG', quality=95)
                            else: tmp_img.convert('RGB').save(temp_img_path, 'JPEG', quality=95)
                            
                            w, h = Image.open(temp_img_path).size
                            calc_width = w / 300.0
                            target_width = Inches(6.0) if calc_width > 6.0 else (Inches(0.5) if calc_width < 0.5 else Inches(calc_width))
                        doc.add_picture(temp_img_path, width=target_width)
                        img_success = True
                    except: pass
                    finally:
                        if os.path.exists(temp_img_path): os.remove(temp_img_path)
                    if img_success and not is_math: line_messages.append(ImageSendMessage(original_content_url=img_url, preview_image_url=img_url))
            except: pass

            if not img_success:
                fallback_txt = clean_math_fallback(raw_content) if is_math else "[" + raw_content + "]"
                if fallback_txt: doc.add_paragraph(fallback_txt)

    if current_text_block.strip():
        cleaned = clean_line_text(current_text_block)
        if cleaned: doc.add_paragraph(cleaned)

    doc_name = "AI_Report_" + uuid.uuid4().hex[:8] + ".docx"
    doc_path = os.path.join("static", doc_name)
    doc.save(doc_path)
    doc_url = SPACE_URL + "/static/" + doc_name + "?openExternalBrowser=1"
    
    return [TextSendMessage(text="📝 解析結果較長或包含圖表，已封裝為 Word 檔供您下載：\n" + doc_url)], doc_url

# ==========================================
# 👑 核心 System Prompt
# ==========================================
GLOBAL_SYSTEM_PROMPT = r"""你是一個極度聰明、邏輯嚴密的 AI 助理，請使用繁體中文回答。

【超強視覺化與圖表指令】
若需要展示數據、流程圖、上網找圖，或複雜數學大矩陣，【絕對必須】使用以下標籤包裝，系統會自動在背景將其渲染為高畫質實體圖片，並排版進 Word 文檔！

1. 【大型數學與矩陣】(例如多層矩陣、大片微積分)：
   必須使用標準 LaTeX 區塊語法，例如 \[ ... \] 或 \begin{bmatrix} ... \end{bmatrix}。系統會自動抓取並渲染圖片。
   但若是【行內簡單算式】(如 a^2+b^2=c^2，分數 1/2)，請直接使用純文字，絕對不要用 LaTeX 標籤包裝。

2. 【圓餅圖/長條圖/折線圖】：
   __CHART__ {"type":"bar","data":{"labels":["A","B"],"datasets":[{"label":"Data","data":[10,20]}]}} __CHART_END__

3. 【流程圖/結構圖】：
   __GRAPH__ digraph G { A -> B; B -> C; } __GRAPH_END__

【最高記憶指令】
你具有完整的長期記憶！仔細閱讀【近期歷史】中用戶上傳過的圖片或對話紀錄並回答，絕對不要回答沒有記憶！

🌟 視覺工具呼叫判斷：
1. 提到「影片、動態、短片」，【必須】輸出 JSON 呼叫影片生成 (tool: "video_generator")。
2. 提到「圖片、照片、畫圖」，【必須】輸出 JSON 呼叫圖片生成 (tool: "image_generator")。
⚠️ JSON 格式 (嚴格獨立輸出)：
__JSON_START__
{"tool": "image_generator", "engine": "safe", "prompt": "[詳細 English Prompt]"}
__JSON_END__"""

# ==========================================
# FastAPI 路由與 Webhook 接收
# ==========================================
@app.get("/")
def read_root(): return {"status": "Ultimate LINE Engine (Pure Native Code) is Running!"}

@app.post("/")
async def callback(request: Request, background_tasks: BackgroundTasks):
    signature = request.headers.get("X-Line-Signature")
    body = await request.body()
    try: background_tasks.add_task(handler.handle, body.decode("utf-8"), signature)
    except InvalidSignatureError: raise HTTPException(status_code=400, detail="Invalid signature")
    return "OK"

# ==========================================
# 核心模組 1：處理文字訊息
# ==========================================
def process_message_background(user_id, user_message, event_timestamp, reply_token):
    messages_to_send = []
    try:
        if user_message == "系統檢查":
            send_line_messages(user_id, reply_token, [TextSendMessage(text="系統排查報告生成中，請點擊下載：\n" + run_system_diagnostic_to_file())])
            return
        if user_message == "⚙️ 大腦設定":
            send_engine_family_menu(user_id, reply_token)
            return
        if user_message.startswith("/engine_family "):
            family = user_message.split(" ")[1]
            send_specific_engine_menu(user_id, reply_token, family)
            return
        if user_message == "🎨 生圖設定":
            send_image_engine_menu(user_id, reply_token)
            return

        if user_message in ["/clear", "清除記憶"]:
            chat_collection.delete_many({"user_id": user_id})
            send_line_messages(user_id, reply_token, [TextSendMessage(text="🧹 記憶已徹底清除！")])
            return

        if user_message == "/toggle_bypass":
            try: pref = prefs_collection.find_one({"user_id": user_id}) or {}
            except: pref = {}
            new_state = not pref.get("bypass_mode", False)
            try: prefs_collection.update_one({"user_id": user_id}, {"$set": {"bypass_mode": new_state}}, upsert=True)
            except: pass
            status = "🟢開啟" if new_state else "🔴關閉"
            send_line_messages(user_id, reply_token, [TextSendMessage(text="⚡ 直通模式已 " + status + "！")])
            return

        txt_match = re.match(r'^/engine\s+(.+)', user_message, re.IGNORECASE)
        if txt_match:
            target = txt_match.group(1).strip()
            try: prefs_collection.update_one({"user_id": user_id}, {"$set": {"text_engine": target, "bypass_mode": False}}, upsert=True)
            except: pass
            send_line_messages(user_id, reply_token, [TextSendMessage(text="⚙️ 思考大腦已切換為：" + TEXT_ENGINES.get(target, target))])
            return

        img_match = re.match(r'^/image_engine\s+(auto|sdxl|anything|z-image|flux|anime|realism|3d|turbo|dark|midjourney|dalle)', user_message, re.IGNORECASE)
        if img_match:
            target = img_match.group(1).lower()
            try: prefs_collection.update_one({"user_id": user_id}, {"$set": {"image_engine": target}}, upsert=True)
            except: pass
            send_line_messages(user_id, reply_token, [TextSendMessage(text="⚙️ 生圖引擎已鎖定為：" + ENGINE_LABELS.get(target, target))])
            return

        user_last_prompt_data = user_last_prompt.get(user_id)
        batch_match = re.match(r'^/batch\s+(\d+)', user_message, re.IGNORECASE)
        if batch_match:
            count = min(int(batch_match.group(1)), 10)
            if not user_last_prompt_data:
                send_line_messages(user_id, reply_token, [TextSendMessage(text="❌ 找不到靈感記憶，請重新叫我畫一張。")])
                return
                
            english_prompt = user_last_prompt_data["prompt"]
            selected_engine = user_last_prompt_data["engine"]
            
            for i in range(count):
                try:
                    seed = uuid.uuid4().hex[:6]
                    image_url = generate_image_and_save(selected_engine, english_prompt, seed)
                    messages_to_send.append(ImageSendMessage(original_content_url=image_url, preview_image_url=image_url))
                except Exception as e: 
                    messages_to_send.append(TextSendMessage(text="❌ 生圖失敗: " + str(e)[:50]))
                time.sleep(1) 
            messages_to_send.append(get_batch_template())
            send_line_messages(user_id, reply_token, messages_to_send)
            return

        try: pref = prefs_collection.find_one({"user_id": user_id}) or {}
        except: pref = {}
        current_text_engine = pref.get("text_engine", "merge")
        current_img_engine = pref.get("image_engine", "auto")
        bypass_mode = pref.get("bypass_mode", False)

        if bypass_mode:
            draw_tool = "video_generator" if any(k in user_message.lower() for k in ["影片", "動態", "短片", "動畫"]) else "image_generator"
            ai_reply = "__JSON_START__\n{\"tool\": \"" + draw_tool + "\", \"engine\": \"" + current_img_engine + "\", \"prompt\": \"" + user_message + "\"}\n__JSON_END__"
            engine_used = "Bypass (無腦直通)"
        else:
            try:
                chat_collection.insert_one({"user_id": user_id, "role": "user", "content": user_message, "timestamp": event_timestamp})
                history_cursor = chat_collection.find({"user_id": user_id}).sort("timestamp", -1).limit(30)
                recent_history = list(history_cursor)[::-1]
            except: recent_history = []

            ai_reply = ""
            error_notice = ""
            engine_used = current_text_engine
            
            hf_messages = [{"role": "system", "content": GLOBAL_SYSTEM_PROMPT}]
            gemini_contents = [GLOBAL_SYSTEM_PROMPT]
            
            for msg in recent_history:
                clean_hist = clean_history_text(msg['content'])
                hf_messages.append({"role": msg["role"], "content": clean_hist[:1500]})
                gemini_contents.append(("User" if msg['role']=='user' else "AI") + ": " + clean_hist)
                
            if current_text_engine == "merge":
                def run_qwen_merge(): return run_hf_model("Qwen/Qwen2.5-7B-Instruct", hf_messages)
                def run_gemini_merge(): return run_gemini(gemini_contents, 'gemini-1.5-flash')

                with concurrent.futures.ThreadPoolExecutor() as executor:
                    fg = executor.submit(run_gemini_merge)
                    fgpt = executor.submit(run_qwen_merge)
                    try: gr = fg.result(timeout=15)
                    except Exception as e: gr = None; err_gr = str(e)
                    try: gptr = fgpt.result(timeout=15)
                    except Exception as e: gptr = None; err_gptr = str(e)
                
                if gr and gptr:
                    merge_prompt = "草稿一：\n" + str(gr) + "\n草稿二：\n" + str(gptr) + "\n請完美整合以上解答。若有 \\[ \\], __CHART__ 等標籤請原封不動保留！"
                    try: 
                        merged_ans = run_gemini([merge_prompt], 'gemini-1.5-flash')
                        ai_reply = "🌟 【雙開整合解答】\n\n" + str(merged_ans)
                    except Exception as merge_e:
                        raw_err = traceback.format_exc()
                        err_url = save_error_log(raw_err)
                        error_notice = "⚠️ [引擎崩潰警告]\n雙開整合過程發生錯誤。\n📝 錯誤日誌：\n" + err_url + "\n\n🔄 已降級為 Gemini 獨立解答...\n\n"
                        ai_reply = gr
                elif gr: 
                    error_notice = "⚠️ [引擎崩潰警告]\n開源引擎已陣亡。\n🔄 已降級為 Gemini 獨立解答...\n\n"
                    ai_reply = gr
                elif gptr: 
                    error_notice = "⚠️ [引擎崩潰警告]\nGemini 引擎已陣亡。\n🔄 已降級為開源引擎獨立解答...\n\n"
                    ai_reply = gptr
                else: 
                    err_url = save_error_log("雙引擎崩潰。\nGemini Err: " + str(locals().get('err_gr')) + "\nQwen Err: " + str(locals().get('err_gptr')))
                    ai_reply = "⚠️ 雙引擎皆超時崩潰。\n📝 錯誤日誌：\n" + err_url

            else:
                try: 
                    ai_reply = smart_engine_runner(current_text_engine, hf_messages, gemini_contents)
                except Exception as e:
                    raw_err = traceback.format_exc()
                    err_url = save_error_log(raw_err)
                    
                    if "402" in str(e) or "401" in str(e) or "403" in str(e):
                        error_notice = "⚠️ [授權受限]\n官方伺服器要求付費 (402) 或需要手動授權。\n📝 日誌已打包：\n" + err_url + "\n\n🔄 自動切換至免費 Gemini 代答...\n\n"
                    elif "DNS" in str(e) or "ConnectionError" in str(e):
                        error_notice = "⚠️ [網路斷線]\n雲端伺服器 DNS 解析失敗。如果你常看到這個，請將此機器人遷移至 Render 或 Zeabur 等穩定的雲端平台！\n📝 日誌已打包：\n" + err_url + "\n\n🔄 自動切換至 Gemini 代答...\n\n"
                    else:
                        error_notice = "⚠️ [引擎崩潰警告]\n您指定的引擎 (" + current_text_engine + ") 已陣亡。\n📝 錯誤日誌已打包：\n" + err_url + "\n\n🔄 已自動為您切換至 Gemini 引擎代答...\n\n"
                    
                    try:
                        ai_reply = run_gemini(gemini_contents, 'gemini-1.5-flash')
                        engine_used = "gemini-1.5-flash (備援)"
                    except Exception as backup_e:
                        backup_err = traceback.format_exc()
                        backup_url = save_error_log(backup_err)
                        error_notice += "⚠️ [備援亦崩潰]\nGemini 也失效，日誌：" + backup_url + "\n\n"
                        ai_reply = "❌ 所有備援引擎全數陣亡！"

            ai_reply = error_notice + str(ai_reply) if ai_reply else error_notice + "⚠️ AI 引擎回傳空值"

        if "__JSON_START__" in ai_reply:
            try:
                match = re.search(r'__JSON_START__(.*?)__JSON_END__', ai_reply, re.DOTALL)
                if match:
                    tool_json = json.loads(match.group(1).strip())
                    if "video" in tool_json.get("tool", ""):
                        messages_to_send.append(TextSendMessage(text="✨ 真影片功能開發中..."))
                    else:
                        img_url = generate_image_and_save(tool_json.get("engine", "auto"), tool_json.get("prompt"), uuid.uuid4().hex[:6])
                        messages_to_send.append(ImageSendMessage(original_content_url=img_url, preview_image_url=img_url))
                        user_last_prompt[user_id] = {"prompt": tool_json.get("prompt", ""), "engine": tool_json.get("engine", "auto")}
            except Exception as e:
                err_url = save_error_log(traceback.format_exc())
                messages_to_send.append(TextSendMessage(text="❌ 生圖執行失敗。\n📝 錯誤日誌：" + err_url))
            
            messages_to_send.append(get_batch_template())
            send_line_messages(user_id, reply_token, messages_to_send)
        else:
            if current_text_engine != engine_used and current_text_engine != "merge" and "⚠️" not in ai_reply: 
                ai_reply = "🔄 [" + TEXT_ENGINES.get(engine_used, engine_used) + "]\n" + ai_reply
            try: chat_collection.insert_one({"user_id": user_id, "role": "assistant", "content": ai_reply, "timestamp": event_timestamp + 1})
            except: pass

            needs_docx = len(ai_reply) > 800 or any(tag in ai_reply for tag in ['__CHART__', '__GRAPH__', '__IMAGE__', '\\[', '\\begin'])
            if needs_docx:
                line_msgs, doc_url = parse_and_build_content(ai_reply, user_id)
                send_line_messages(user_id, reply_token, line_msgs)
            else:
                ai_reply = clean_line_text(process_inline_math(ai_reply))
                send_line_messages(user_id, reply_token, [TextSendMessage(text=ai_reply)])
                
    except Exception as fatal_e:
        err_url = save_error_log(traceback.format_exc())
        send_line_messages(user_id, reply_token, [TextSendMessage(text="❌ 核心運算崩潰。\n📝 詳細錯誤日誌請下載查看：\n" + err_url)])

@handler.add(MessageEvent, message=TextMessage)
def handle_text_event(event):
    show_loading_animation(event.source.user_id)
    threading.Thread(target=process_message_background, args=(event.source.user_id, event.message.text.strip(), event.timestamp, event.reply_token)).start()

# ==========================================
# 核心模組 2：接收圖片訊息
# ==========================================
def process_image_background(user_id, message_id, event_timestamp, reply_token):
    try:
        message_content = line_bot_api.get_message_content(message_id)
        image_bytes = b""
        for chunk in message_content.iter_content(): image_bytes += chunk
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode != 'RGB': img = img.convert('RGB')
        img.thumbnail((2048, 2048)) 
        path = os.path.join("static", f"{uuid.uuid4().hex}.jpg")
        img.save(path, format="JPEG", quality=95)
        
        try:
            try:
                reply_text = run_gemini([img, "請詳細分析這張圖片的內容，如果是題目請給出詳解。"], 'gemini-1.5-flash')
            except Exception as e:
                err_url = save_error_log(traceback.format_exc())
                reply_text = "⚠️ Gemini API 處理圖片失敗。\n📝 錯誤日誌已打包：" + err_url + "\n目前僅支援 Gemini 分析圖片，請稍後再試。"
            
            reply_text = str(reply_text) if reply_text else "⚠️ AI 引擎拒絕看圖或回傳了空白內容。"

            needs_docx = len(reply_text) > 800 or any(tag in reply_text for tag in ['__CHART__', '__GRAPH__', '__SEARCH__', '\\[', '\\begin'])
            if needs_docx:
                line_messages, doc_url = parse_and_build_content(reply_text, user_id)
                send_line_messages(user_id, reply_token, line_messages)
            else:
                send_line_messages(user_id, reply_token, [TextSendMessage(text=clean_line_text(process_inline_math(reply_text)))])
            
            try:
                b64_img = image_to_base64(path)
                chat_collection.insert_one({"user_id": user_id, "role": "user", "content": "[上傳了圖片]", "image_b64": b64_img, "timestamp": event_timestamp})
                chat_collection.insert_one({"user_id": user_id, "role": "assistant", "content": reply_text, "timestamp": event_timestamp+1})
            except: pass
        except Exception as api_e:
            err_url = save_error_log(traceback.format_exc())
            send_line_messages(user_id, reply_token, [TextSendMessage(text="❌ 圖片分析失敗。\n📝 錯誤日誌：\n" + err_url)])
        finally:
            if os.path.exists(path): os.remove(path)
    except Exception as e:
        err_url = save_error_log(traceback.format_exc())
        send_line_messages(user_id, reply_token, [TextSendMessage(text="❌ 接收圖片失敗。\n📝 錯誤日誌：\n" + err_url)])

@handler.add(MessageEvent, message=ImageMessage)
def handle_image_message(event):
    show_loading_animation(event.source.user_id)
    threading.Thread(target=process_image_background, args=(event.source.user_id, event.message.id, event.timestamp, event.reply_token)).start()

# ==========================================
# 核心模組 3：處理檔案訊息
# ==========================================
def process_file_background(user_id, file_name, message_id, event_timestamp, reply_token):
    local_pdf_path = os.path.join("static", f"{uuid.uuid4().hex}.pdf")
    try:
        message_content = line_bot_api.get_message_content(message_id)
        with open(local_pdf_path, "wb") as f:
            for chunk in message_content.iter_content(): f.write(chunk)
        
        reader = PdfReader(local_pdf_path)
        ext_txt = "".join([p.extract_text() or "" for p in reader.pages])
                
        prompt = "請幫我全面分析這份 PDF，用繁體中文回覆。\n⚠️ 若有數學式，請使用 LaTeX 區塊包裝在 \\[ ... \\] 內以便渲染。"
        ai_reply = ""
        error_notice = ""
        try: pref = prefs_collection.find_one({"user_id": user_id}) or {}
        except: pref = {}
        current_text_engine = pref.get("text_engine", "merge")

        try:
            if current_text_engine == "merge":
                def run_qwen_pdf(): return run_hf_model("Qwen/Qwen2.5-7B-Instruct", [{"role": "user", "content": "PDF內容：\n" + ext_txt[:6000] + "\n\n回答：" + prompt}])
                def run_gemini_pdf(): return run_gemini(["PDF內容：\n" + ext_txt[:15000] + "\n\n回答：" + prompt], 'gemini-1.5-flash')
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    fg = executor.submit(run_gemini_pdf)
                    fgpt = executor.submit(run_qwen_pdf)
                    try: gr = fg.result(timeout=25)
                    except Exception as e: gr = None; err_gr = str(e)
                    try: gptr = fgpt.result(timeout=20)
                    except Exception as e: gptr = None; err_gptr = str(e)
                
                if gr and gptr:
                    merge_prompt = "草稿一：\n" + str(gr) + "\n草稿二：\n" + str(gptr) + "\n請完美整合。若有 \\[ \\], __CHART__ 等標籤請原封不動保留！"
                    try: 
                        merged_ans = run_gemini([merge_prompt], 'gemini-1.5-flash')
                        ai_reply = "🌟 【雙引擎終極整合解答】\n\n" + str(merged_ans)
                    except Exception as e: 
                        err_url = save_error_log(traceback.format_exc())
                        ai_reply = "⚠️ [警告：整合過程超載，已降級為 Gemini 獨立解答]\n📝 錯誤日誌已打包：" + err_url + "\n\n" + str(gr)
                elif gr: 
                    ai_reply = "⚠️ [警告：Qwen 引擎連線失敗，本次合併已降級為 Gemini 獨立解答]\n\n" + str(gr)
                elif gptr: 
                    ai_reply = "⚠️ [警告：Gemini 引擎連線失敗，本次合併已降級為 Qwen 獨立解答]\n\n" + str(gptr)
                else: 
                    err_url = save_error_log("雙引擎崩潰。\nGemini Err: " + str(locals().get('err_gr')) + "\nQwen Err: " + str(locals().get('err_gptr')))
                    ai_reply = "⚠️ 雙引擎皆崩潰。\n📝 錯誤日誌已打包：" + err_url

            else:
                try: 
                    if current_text_engine.startswith("Qwen") or current_text_engine.startswith("meta-llama") or current_text_engine.startswith("mistralai"):
                        ai_reply = run_hf_model(current_text_engine, [{"role": "user", "content": "PDF內容：\n" + ext_txt[:6000] + "\n\n回答：" + prompt}])
                    else:
                        ai_reply = run_gemini(["PDF內容：\n" + ext_txt[:15000] + "\n\n回答：" + prompt], current_text_engine)
                except Exception as e:
                    raw_err = traceback.format_exc()
                    err_url = save_error_log(raw_err)
                    error_notice = "⚠️ [引擎崩潰警告]\n您選擇的引擎處理 PDF 失敗。\n📝 錯誤日誌已打包：\n" + err_url + "\n\n🔄 自動為您切換至 Gemini 引擎代答...\n\n"
                    try:
                        ai_reply = run_gemini(["PDF內容：\n" + ext_txt[:15000] + "\n\n回答：" + prompt], 'gemini-1.5-flash')
                    except Exception as backup_e:
                        backup_err = save_error_log(traceback.format_exc())
                        ai_reply = "❌ 備援 Gemini 亦處理失敗。\n📝 錯誤日誌：" + backup_err
                        
            ai_reply = error_notice + str(ai_reply)

        except Exception as e:
            err_url = save_error_log(traceback.format_exc())
            send_line_messages(user_id, reply_token, [TextSendMessage(text="❌ PDF 解析發生系統級崩潰。\n📝 錯誤日誌：\n" + err_url)])
            return

        ai_reply = str(ai_reply) if ai_reply else "⚠️ AI 引擎拒絕讀取 PDF 或回傳了空白內容。"

        try:
            chat_collection.insert_one({"user_id": user_id, "role": "user", "content": "[上傳 PDF]: " + file_name, "timestamp": event_timestamp})
            chat_collection.insert_one({"user_id": user_id, "role": "assistant", "content": ai_reply, "timestamp": event_timestamp + 1})
        except: pass

        needs_docx = len(ai_reply) > 800 or any(tag in ai_reply for tag in ['__CHART__', '__GRAPH__', '__IMAGE__', '\\[', '\\begin'])
        if needs_docx:
            line_messages, doc_url = parse_and_build_content(ai_reply, user_id)
            send_line_messages(user_id, reply_token, line_messages)
        else:
            send_line_messages(user_id, reply_token, [TextSendMessage(text=clean_line_text(process_inline_math(ai_reply)))])
        
    except Exception as fatal_e:
        err_url = save_error_log(traceback.format_exc())
        send_line_messages(user_id, reply_token, [TextSendMessage(text="❌ 處理 PDF 檔案發生未預期崩潰。\n📝 錯誤日誌請下載查看：\n" + err_url)])

@handler.add(MessageEvent, message=FileMessage)
def handle_file_event(event):
    user_id = event.source.user_id
    file_name = event.message.file_name
    if not file_name.lower().endswith('.pdf'):
        line_bot_api.reply_message(event.reply_token, TextSendMessage(text="⚠️ 目前只支援 PDF 檔案喔！", quick_reply=get_quick_replies(user_id)))
        return
    show_loading_animation(user_id)
    threading.Thread(target=process_file_background, args=(user_id, file_name, event.message.id, event.timestamp, event.reply_token)).start()
